"""生成 pandoc reference.docx，预置接近深大本科毕设格式的样式。

pandoc 会用该 docx 中的命名样式渲染正文、标题、图表说明等。这里给出
一个近似实现：
- 正文：宋体，五号（10.5pt），段前段后 0.5 行，单倍行距，首行缩进 2 字符
- 一级标题：黑体，三号（16pt），加粗
- 二级标题：黑体，小三号（15pt），加粗
- 三级标题：黑体，四号（14pt），加粗
- 图表说明：宋体，小五号（9pt），居中

注意：docx 的"小号"与磅值不完全一一对应，这里采用常见教务文档
实现，生成后请再以深大官方模板对照修订。
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

BUILD_DIR = Path(__file__).resolve().parent.parent / "build"
REFERENCE_DOCX = BUILD_DIR / "reference.docx"


def _set_font(run, *, eastasia: str, western: str, size: float, bold: bool = False) -> None:
    run.font.name = western
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), eastasia)
    rFonts.set(qn("w:ascii"), western)
    rFonts.set(qn("w:hAnsi"), western)


def _style(doc: Document, name: str, *, eastasia: str, western: str,
           size: float, bold: bool = False,
           alignment: int = WD_ALIGN_PARAGRAPH.JUSTIFY,
           first_line_indent: float = 0.0) -> None:
    try:
        style = doc.styles[name]
    except KeyError:
        return
    style.font.name = western
    style.font.size = Pt(size)
    style.font.bold = bold
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), eastasia)
    rFonts.set(qn("w:ascii"), western)
    rFonts.set(qn("w:hAnsi"), western)
    pf = style.paragraph_format
    pf.alignment = alignment
    pf.space_before = Pt(size * 0.5)
    pf.space_after = Pt(size * 0.5)
    if first_line_indent:
        pf.first_line_indent = Pt(size * first_line_indent)


def main() -> None:
    BUILD_DIR.mkdir(exist_ok=True)
    doc = Document()

    # 正文
    _style(doc, "Normal", eastasia="宋体", western="Times New Roman",
           size=10.5, bold=False,
           alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=2)

    # 各级标题
    _style(doc, "Heading 1", eastasia="黑体", western="Times New Roman",
           size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    _style(doc, "Heading 2", eastasia="黑体", western="Times New Roman",
           size=15, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    _style(doc, "Heading 3", eastasia="黑体", western="Times New Roman",
           size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # 图表说明
    _style(doc, "Caption", eastasia="宋体", western="Times New Roman",
           size=9, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 题名（用于中文题名位置）
    _style(doc, "Title", eastasia="华文中宋", western="Times New Roman",
           size=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 默认页脚：居中 PAGE 字段（页码自动）
    section = doc.sections[0]
    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 清空已有文本
    for r in list(footer_para.runs):
        r._element.getparent().remove(r._element)
    # 插入 PAGE 字段
    run = footer_para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_end)

    doc.save(REFERENCE_DOCX)
    print(f"[reference.docx] -> {REFERENCE_DOCX}")


if __name__ == "__main__":
    main()

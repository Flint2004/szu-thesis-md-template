"""DDSVM 毕业论文构建脚本。

工作流：
1. 读取 thesis/ 下各章 markdown
2. 对正文章节做引用重写：把 [N] / [N-M] / [N, M] 改写为 pandoc 超链接
   `[\\[N\\]](#ref_N)`，对应参考文献章节中后处理插入的 bookmark `ref_N`
3. 合并所有章节为单一 markdown，调用 pandoc 转 docx
4. 用 python-docx 后处理生成的 docx，套用深大本科毕设格式：
   - 在「目录」标题段后插入 Word TOC 字段
   - 「目录」标题降级为非 Heading（不出现在自身目录里）
   - 摘要/Abstract 引导段落按"五号楷体 + 标识符小四加粗"渲染
   - 章节标题统一为黑色、宋体（majorEastAsia 主题）、首行无缩进
   - 引用上标化：[N] 显示为右上标
   - 表格按 booktabs 风格美化：整表居中、表头加粗、三道横线
   - 图/表 caption 段落以及含图片的段落统一居中
   - 参考文献条目段落补插 bookmark
   - OMML 中纯数字 run 设为 plain，避免数字被 Word 当变量斜体化
   - 元素式内部超链接转为 HYPERLINK 字段码（Mac Word 兼容性）
"""

from __future__ import annotations

import re
from pathlib import Path

import pypandoc
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
THESIS_DIR = ROOT / "thesis"
FIGURES_DIR = ROOT / "figures"
BUILD_DIR = ROOT / "build"
REFERENCE_DOCX = BUILD_DIR / "reference.docx"
OUTPUT_DOCX = BUILD_DIR / "thesis.docx"
MERGED_MD = BUILD_DIR / "thesis.md"

CHAPTERS = [
    "00_frontmatter.md",
    "01_intro.md",
    "02_related.md",
    "03_method.md",
    "04_theory.md",
    "05_experiments.md",
    "06_conclusion.md",
    "07_references.md",
    "08_ack.md",
    "09_abstract_en.md",
    "10_appendix.md",
]

# ---------------------------------------------------------------------------
# 命名空间常量
# ---------------------------------------------------------------------------
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


# ===========================================================================
# 阶段 1 / 2：合并 + 引用重写
# ===========================================================================

CITATION_RE = re.compile(r"\[([\d\s,\-\u2013]+)\]")


def rewrite_citations(text: str) -> str:
    """把 [N] / [N-M] / [N, M] 改写为 pandoc 超链接，并整体包成上标。

    用 pandoc superscript 语法 `^...^` 包住整组引用，确保方括号、逗号
    与数字一同上标。superscript 内空格用 `\\ ` 转义（pandoc 要求）。
    """

    def repl(m: re.Match[str]) -> str:
        inner = m.group(1).strip()
        # 跳过非数字开头或以 0 开头（数学区间 [0,1] 等）
        if not inner or not inner[0].isdigit() or inner[0] == "0":
            return m.group(0)
        parts = [p.strip() for p in inner.split(",")]
        if len(parts) == 1:
            seg = parts[0]
            m2 = re.match(r"^(\d+)[\-\u2013](\d+)$", seg)
            if m2:
                return rf"^[\[{seg}\]](#ref_{m2.group(1)})^"
            if seg.isdigit():
                return rf"^[\[{seg}\]](#ref_{seg})^"
            return m.group(0)
        linked: list[str] = []
        for seg in parts:
            m2 = re.match(r"^(\d+)[\-\u2013](\d+)$", seg)
            if m2:
                linked.append(rf"[{seg}](#ref_{m2.group(1)})")
            elif seg.isdigit():
                linked.append(rf"[{seg}](#ref_{seg})")
            else:
                return m.group(0)
        joined = r",\ ".join(linked)  # \  保留空格但保持 superscript 不断开
        return rf"^\[{joined}\]^"

    return CITATION_RE.sub(repl, text)


ESCAPED_QUOTE_RE = re.compile(r"\\([\"'])")


def _unescape_quotes(text: str) -> str:
    """把 `\\"` / `\\'` 还原为 `"` / `'`。

    作者在 md 源中写 `\\"` 以避免 fix_quotes.py 的自动转中文引号；到 docx
    阶段应是字面直引号。注意保护：反斜杠本身若需出现（如 `\\\\`）应写两
    次，此 regex 只吃紧跟 `"`/`'` 的单反斜杠。
    """
    return ESCAPED_QUOTE_RE.sub(r"\1", text)


def merge_markdown() -> Path:
    BUILD_DIR.mkdir(exist_ok=True)
    parts: list[str] = []
    for name in CHAPTERS:
        fp = THESIS_DIR / name
        raw = fp.read_text(encoding="utf-8")
        # 参考文献章节自身不重写（避免每条开头的 [N] 被改）
        if name == "07_references.md":
            body = raw
        else:
            body = rewrite_citations(raw)
        # 最后一步：把 `\"` / `\'` 还原为直引号（脚本保留规则的配对操作）
        body = _unescape_quotes(body)
        parts.append(body)
    MERGED_MD.write_text("\n\n".join(parts), encoding="utf-8")
    return MERGED_MD


# ===========================================================================
# 阶段 3：调用 pandoc 转 docx
# ===========================================================================


def convert_to_docx(merged_md: Path) -> Path:
    extra_args = [
        f"--resource-path={FIGURES_DIR}:{THESIS_DIR}:{ROOT}",
        "--standalone",
    ]
    if REFERENCE_DOCX.exists():
        extra_args.append(f"--reference-doc={REFERENCE_DOCX}")
    pypandoc.convert_file(
        str(merged_md),
        "docx",
        # -smart 关闭 pandoc 的"智能引号"扩展，避免把参考文献等处残留的
        # 英文直引号 "..." 自动转为 Unicode 弯引号 "..."（恰好与中文引号同码）
        format="markdown-smart+tex_math_dollars+raw_tex+fenced_divs",
        outputfile=str(OUTPUT_DOCX),
        extra_args=extra_args,
    )
    return OUTPUT_DOCX


# ===========================================================================
# 阶段 4：python-docx + lxml 后处理
# ===========================================================================

REF_LINE_RE = re.compile(r"^\[(\d+)\]")
DIGIT_RE = re.compile(r"^[0-9.,\s]+$")
ABSTRACT_MARK_CN_RE = re.compile(r"^【(摘要|关键词)】")
ABSTRACT_MARK_EN_RE = re.compile(r"^【(Abstract|Key words)】")
CHAPTER_HEAD_RE = re.compile(r"^第\d+章\s")
SUBSEC_HEAD_RE = re.compile(r"^\d+\.\d+\s")  # "1.1 ", "2.3 " ...
SUBSUBSEC_HEAD_RE = re.compile(r"^\d+\.\d+\.\d+\s")  # "1.1.1 ", "3.5.2 " ...
FIG_CAPTION_PREFIX_RE = re.compile(r"^(图\s*\d+(?:\.\d+)?)")
TBL_CAPTION_PREFIX_RE = re.compile(r"^(表\s*\d+(?:\.\d+)?)")
HIDE_HEADING_TEXTS = {"【摘要】", "【Abstract】"}

# 图片宽度上限（EMU）；约 3.6 英寸 / 9.14 cm
DEFAULT_MAX_PIC_CX = 3291840


def _make_toc_paragraph():
    """构造一个带 Word TOC 字段的 <w:p> 元素，only-抓 Heading 1-3。"""
    p = OxmlElement("w:p")

    def add_run(children: list[OxmlElement]) -> None:
        r = OxmlElement("w:r")
        for c in children:
            r.append(c)
        p.append(r)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    add_run([begin])

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-3" \h \z \u '
    add_run([instr])

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    add_run([sep])

    t = OxmlElement("w:t")
    t.text = "请在 Word 中右键此段并选择\u201c更新域\u201d以生成目录。"
    add_run([t])

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    add_run([end])

    return p


def _add_bookmark(para, name: str, bm_id: int) -> None:
    """在段落开头与结尾各插入 bookmarkStart/End。"""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bm_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bm_id))
    pPr = para._element.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(start)
    else:
        para._element.insert(0, start)
    para._element.append(end)


def _paragraph_has_image(para) -> bool:
    return para._element.find(".//" + qn("w:drawing")) is not None


def _para_set_no_indent(para) -> None:
    """段落级清零左缩进与首行缩进。"""
    pPr = para._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        para._element.insert(0, pPr)
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    for k in ("w:left", "w:leftChars", "w:firstLine", "w:firstLineChars"):
        ind.set(qn(k), "0")


def _run_set_font(run, *, eastasia: str | None = None, ascii_: str | None = None,
                  size_pt: float | None = None, color: str | None = None,
                  bold: bool | None = None, superscript: bool = False) -> None:
    """对单个 w:r 元素设置字体/字号/颜色/粗体/上标。"""
    rPr = run.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        run.insert(0, rPr)
    if eastasia or ascii_:
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        if eastasia:
            rFonts.set(qn("w:eastAsia"), eastasia)
            rFonts.set(qn("w:hint"), "eastAsia")
        if ascii_:
            rFonts.set(qn("w:ascii"), ascii_)
            rFonts.set(qn("w:hAnsi"), ascii_)
            rFonts.set(qn("w:cs"), ascii_)
    if size_pt is not None:
        v = str(int(round(size_pt * 2)))  # half-points
        for tag in ("w:sz", "w:szCs"):
            el = rPr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                rPr.append(el)
            el.set(qn("w:val"), v)
    if color is not None:
        col = rPr.find(qn("w:color"))
        if col is None:
            col = OxmlElement("w:color")
            rPr.append(col)
        col.set(qn("w:val"), color)
    if bold is not None:
        for tag in ("w:b", "w:bCs"):
            el = rPr.find(qn(tag))
            if el is None and bold:
                rPr.append(OxmlElement(tag))
            elif el is not None and not bold:
                rPr.remove(el)
    if superscript:
        va = rPr.find(qn("w:vertAlign"))
        if va is None:
            va = OxmlElement("w:vertAlign")
            rPr.append(va)
        va.set(qn("w:val"), "superscript")


# ---------------------------------------------------------------------------
# 4.1 章节标题：去缩进、改字体、改颜色（黑色）
# ---------------------------------------------------------------------------


def _restyle_chapter_headings(doc: Document) -> tuple[int, int, int]:
    """统一处理 Heading 1/2/3 样式，返回 (h1_count, h2_count, h3_count)。

    - Heading 1：宋体 + Times + 黑色 + 无缩进（章标题或【摘要】等）
    - Heading 2 节标题（"1.1 ..."）：宋体 + Times + 黑色 + 加粗 + 15pt（小三）+ 无缩进
    - Heading 3 节标题（"1.1.1 ..."）：宋体 + Times + 黑色 + 加粗 + 14pt（四号）+ 无缩进
    - 其余 Heading：仅强制黑色
    """
    h1 = h2 = h3 = 0
    for para in doc.paragraphs:
        sn = para.style.name
        if not sn.startswith("Heading"):
            continue
        text = para.text.strip()

        if sn == "Heading 1":
            _para_set_no_indent(para)
            for r in para._element.findall(qn("w:r")):
                _run_set_font(r, eastasia="黑体", ascii_="Times New Roman",
                              color="000000")
            h1 += 1
        elif sn == "Heading 2" and SUBSEC_HEAD_RE.match(text):
            _para_set_no_indent(para)
            for r in para._element.findall(qn("w:r")):
                _run_set_font(r, eastasia="黑体", ascii_="Times New Roman",
                              size_pt=15, bold=True, color="000000")
            h2 += 1
        elif sn == "Heading 3" and SUBSUBSEC_HEAD_RE.match(text):
            _para_set_no_indent(para)
            for r in para._element.findall(qn("w:r")):
                _run_set_font(r, eastasia="黑体", ascii_="Times New Roman",
                              size_pt=14, bold=True, color="000000")
            h3 += 1
        else:
            for r in para._element.findall(qn("w:r")):
                _run_set_font(r, color="000000")
    return h1, h2, h3


def _restyle_special_headings(doc: Document) -> tuple[int, int]:
    """对「参考文献」与「致谢」两个 H1 套深大专用样式。

    - 参考文献：五号（10.5pt）楷体加粗，顶格
    - 致谢：小四号（12pt）黑体加粗，居中
    (两个都不覆盖通用 H1 的 firstLine=0 / color / page break。)
    """
    refs = acks = 0
    for para in doc.paragraphs:
        if para.style.name != "Heading 1":
            continue
        t = para.text.strip()
        if t == "参考文献":
            for r in para._element.findall(qn("w:r")):
                _run_set_font(r, eastasia="华文楷体", ascii_="Times New Roman",
                              size_pt=10.5, bold=True, color="000000")
            refs += 1
        elif t == "致谢":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para._element.findall(qn("w:r")):
                _run_set_font(r, eastasia="黑体", ascii_="Times New Roman",
                              size_pt=12, bold=True, color="000000")
            acks += 1
    return refs, acks


def _restyle_reference_entries(doc: Document) -> int:
    """参考文献条目（`[N] ...`）：悬挂缩进 + 段前段后 0.5 行 + 单倍行距 + 小五。

    参数值参考深大最终稿：left=419 hanging=419 hangingChars=233（悬挂 2 字符），
    spacing before/after=181 (beforeLines/afterLines=50 = 0.5 行)，line=240 lineRule=auto。
    """
    fixed = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not REF_LINE_RE.match(text):
            continue
        pPr = para._element.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            para._element.insert(0, pPr)

        # ind：悬挂缩进
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        ind.set(qn("w:left"), "419")
        ind.set(qn("w:leftChars"), "0")
        ind.set(qn("w:hanging"), "419")
        ind.set(qn("w:hangingChars"), "233")
        for k in ("w:firstLine", "w:firstLineChars"):
            if ind.get(qn(k)) is not None:
                del ind.attrib[qn(k)]

        # spacing
        sp = pPr.find(qn("w:spacing"))
        if sp is None:
            sp = OxmlElement("w:spacing")
            pPr.append(sp)
        sp.set(qn("w:before"), "181")
        sp.set(qn("w:beforeLines"), "50")
        sp.set(qn("w:after"), "181")
        sp.set(qn("w:afterLines"), "50")
        sp.set(qn("w:line"), "240")
        sp.set(qn("w:lineRule"), "auto")

        # run 字号小五 (9pt = sz 18)
        for r in para._element.findall(qn("w:r")):
            _run_set_font(r, size_pt=9)

        fixed += 1
    return fixed


def _hide_marker_headings(doc: Document) -> int:
    """把【摘要】/【Abstract】两个 Heading 1 段落"近不可见"化。

    *不能用 vanish*：vanish 段落 Word 在更新 TOC 时会跳过。
    改用：白色字 + 1pt 字号 + 紧凑行距，视觉几乎不可见，但保留 outline level
    让 TOC 仍能抓到并显示"【摘要】"/"【Abstract】"。
    """
    fixed = 0
    for para in doc.paragraphs:
        if para.style.name != "Heading 1":
            continue
        if para.text.strip() not in HIDE_HEADING_TEXTS:
            continue
        # run 级：白字 + 1pt
        for r in para._element.findall(qn("w:r")):
            rPr = r.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                r.insert(0, rPr)
            # 移除可能残留的 vanish
            v = rPr.find(qn("w:vanish"))
            if v is not None:
                rPr.remove(v)
            _run_set_font(r, color="FFFFFF", size_pt=1)
        # 段落级：spacing before/after 0、line exact 20 twentieths-of-pt = 1pt
        pPr = para._element.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            para._element.insert(0, pPr)
        existing = pPr.find(qn("w:spacing"))
        if existing is not None:
            pPr.remove(existing)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:line"), "20")
        spacing.set(qn("w:lineRule"), "exact")
        pPr.append(spacing)
        fixed += 1
    return fixed


def _add_page_breaks_before_h1(doc: Document) -> int:
    """非首个 Heading 1 前插入分页符。

    首个 H1（【摘要】）由 _setup_section_pagenum 的 section-break (nextPage)
    负责换页，避免重复 page break 造成额外空页。
    """
    h1_paras = [p for p in doc.paragraphs if p.style.name == "Heading 1"]
    inserted = 0
    for p in h1_paras[1:]:
        pPr = p._element.find(qn("w:pPr"))
        new_r = OxmlElement("w:r")
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        new_r.append(br)
        if pPr is not None:
            pPr.addnext(new_r)
        else:
            p._element.insert(0, new_r)
        inserted += 1
    return inserted


def _setup_section_pagenum(doc: Document) -> int:
    """前置部分（目录）使用罗马数字页码，正文（从【摘要】起）使用阿拉伯。

    实现：
    - 在【摘要】段之前插入一个 section break 段，其 sectPr 配 lowerRoman、start=1
      （这一段定义 section 1 的属性，即目录所在 section）
    - body 末尾原始 sectPr 加 pgNumType decimal start=1（section 2 = 摘要起）
    """
    body = doc.element.body

    # 取 body 末尾原始 sectPr，复制其 pgSz/pgMar/cols/docGrid 给中间 section break 用
    final_sectPr = body.find(qn("w:sectPr"))
    if final_sectPr is None:
        return 0

    # 确保 final sectPr 含 decimal pgNumType
    pn = final_sectPr.find(qn("w:pgNumType"))
    if pn is None:
        pn = OxmlElement("w:pgNumType")
        final_sectPr.append(pn)
    pn.set(qn("w:fmt"), "decimal")
    pn.set(qn("w:start"), "1")

    # 找【摘要】Heading 1 段
    target = None
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and p.text.strip() in HIDE_HEADING_TEXTS:
            target = p
            break
    if target is None:
        return 0

    # 构造 section break 段：含 sectPr 复制 pgSz/pgMar，加 pgNumType lowerRoman + nextPage
    from copy import deepcopy

    sb_para = OxmlElement("w:p")
    sb_pPr = OxmlElement("w:pPr")
    sb_para.append(sb_pPr)

    sb_sectPr = OxmlElement("w:sectPr")
    # 复制页面尺寸、边距、栏、网格 + 页眉/页脚 reference（让前置 section 也有页脚）
    for child_tag in (
        "w:headerReference", "w:footerReference",
        "w:pgSz", "w:pgMar", "w:cols", "w:docGrid",
    ):
        for existing in final_sectPr.findall(qn(child_tag)):
            sb_sectPr.append(deepcopy(existing))
    sb_pgnum = OxmlElement("w:pgNumType")
    sb_pgnum.set(qn("w:fmt"), "lowerRoman")
    sb_pgnum.set(qn("w:start"), "1")
    sb_sectPr.append(sb_pgnum)
    sb_type = OxmlElement("w:type")
    sb_type.set(qn("w:val"), "nextPage")
    sb_sectPr.append(sb_type)
    sb_pPr.append(sb_sectPr)

    target._element.addprevious(sb_para)
    return 1


def _resize_images(doc: Document, max_cx: int = DEFAULT_MAX_PIC_CX) -> int:
    """所有内联图片 cx 限制不超过 max_cx；超过则等比缩放（含 cy 与 a:ext）。"""
    body = doc.element.body
    A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
    WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    extent_tag = f"{{{WP_NS}}}extent"
    a_ext_tag = f"{{{A_NS}}}ext"

    fixed = 0
    for ext in body.iter(extent_tag):
        cx = int(ext.get("cx", "0"))
        cy = int(ext.get("cy", "0"))
        if cx <= max_cx or cx == 0:
            continue
        ratio = max_cx / cx
        new_cx = max_cx
        new_cy = int(cy * ratio)
        ext.set("cx", str(new_cx))
        ext.set("cy", str(new_cy))
        # 同步 a:ext（位于 pic:spPr/a:xfrm/a:ext 中）
        drawing = ext.getparent().getparent()  # wp:inline -> w:drawing
        for a_ext in drawing.iter(a_ext_tag):
            a_ext.set("cx", str(new_cx))
            a_ext.set("cy", str(new_cy))
        fixed += 1
    return fixed


def _raw_pstyle(para) -> str:
    """直接读 <w:pStyle w:val="..."/>。

    pandoc 可能写出 reference.docx 未定义的样式名（如 ImageCaption），
    python-docx 解析时会 fallback 到 Normal，这里绕开。
    """
    pPr = para._element.find(qn("w:pPr"))
    if pPr is None:
        return ""
    ps = pPr.find(qn("w:pStyle"))
    if ps is None:
        return ""
    return ps.get(qn("w:val"), "")


def _is_real_caption(para) -> str | None:
    """返回 'figure'/'table'/None：

    - figure：pandoc 自动生成的 ImageCaption / Caption / CaptionedFigure 等样式段
    - table：段落紧随其后即 <w:tbl>（即表上方的 caption 段）
    - None：其它（包括正文里"图 3.3 描述了…"这种引用句）
    """
    raw = _raw_pstyle(para)
    if "Caption" in raw or "Caption" in para.style.name:
        return "figure"
    nxt = para._element.getnext()
    if nxt is not None and nxt.tag == qn("w:tbl"):
        return "table"
    return None


def _restyle_captions(doc: Document) -> int:
    """真正的图/表 caption 段：所有 run（含 OMML 公式）字号 9pt 小五 + 居中。

    bold 由 markdown 源决定（表 caption 用 `**表 N.M**` 带粗，图 caption 不带）。
    """
    fixed = 0
    for para in doc.paragraphs:
        if _is_real_caption(para) is None:
            continue
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 段落里所有 w:rPr（普通 run + 公式内嵌 ctrlPr）一并设 sz=18
        for rPr in para._element.iter(qn("w:rPr")):
            sz = rPr.find(qn("w:sz"))
            if sz is None:
                sz = OxmlElement("w:sz")
                rPr.append(sz)
            sz.set(qn("w:val"), "18")
            szCs = rPr.find(qn("w:szCs"))
            if szCs is None:
                szCs = OxmlElement("w:szCs")
                rPr.append(szCs)
            szCs.set(qn("w:val"), "18")
        # 没有 rPr 的 w:r 也补一个含 sz 的 rPr
        for r in para._element.findall(qn("w:r")):
            if r.find(qn("w:rPr")) is None:
                _run_set_font(r, size_pt=9)
        fixed += 1
    return fixed


# ---------------------------------------------------------------------------
# 4.2 摘要 / Abstract 引导段：华文楷体 + 标识符加粗 12pt
# ---------------------------------------------------------------------------


def _split_run_at(run, split_idx: int):
    """把 w:r 在 w:t.text 的第 split_idx 处拆为两个连续 run。

    返回 (前半 run, 后半 run)。两者继承原 run 的 rPr。
    """
    from copy import deepcopy

    t_el = run.find(qn("w:t"))
    if t_el is None or t_el.text is None:
        return run, None
    full = t_el.text
    if split_idx <= 0 or split_idx >= len(full):
        return run, None

    head, tail = full[:split_idx], full[split_idx:]
    t_el.text = head
    t_el.set(qn("xml:space"), "preserve")

    new_run = deepcopy(run)
    new_t = new_run.find(qn("w:t"))
    new_t.text = tail
    new_t.set(qn("xml:space"), "preserve")
    run.addnext(new_run)
    return run, new_run


def _restyle_abstract_paragraphs(doc: Document) -> int:
    """识别以【摘要】/【关键词】/【Abstract】/【Key words】开头的段落。

    - 整段顶格：`ind left=0 leftChars=0 firstLine=0 firstLineChars=0`
    - 中文：整段华文楷体（10.5pt 五号 = 默认）
    - 英文：整段 Times New Roman
    - 标识符部分：独立 run，加粗、字号 12pt（小四）
    - 标识符与正文如果在同一个 run，先拆分再分别上样式
    """
    fixed = 0
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            continue
        text = para.text.strip()
        m_cn = ABSTRACT_MARK_CN_RE.match(text)
        m_en = ABSTRACT_MARK_EN_RE.match(text)
        if not (m_cn or m_en):
            continue

        _para_set_no_indent(para)  # 顶格

        is_cn = bool(m_cn)
        font_ea = "华文楷体" if is_cn else "Times New Roman"
        font_lat = "Times New Roman"
        marker = m_cn.group(0) if m_cn else m_en.group(0)

        runs = list(para._element.findall(qn("w:r")))
        # 先统一字体（不动字号、加粗）
        for r in runs:
            _run_set_font(r, eastasia=font_ea, ascii_=font_lat)

        # 找首个含 marker 的 run，并把 marker 部分拆出独立 run
        for r in runs:
            t_el = r.find(qn("w:t"))
            run_text = t_el.text if t_el is not None and t_el.text else ""
            if marker in run_text:
                idx = run_text.index(marker) + len(marker)
                # 1) 若 marker 不在 run 开头，先把 run 在 marker 起点处切掉前缀
                start = run_text.index(marker)
                if start > 0:
                    _, second = _split_run_at(r, start)
                    if second is None:
                        break
                    r = second  # 后续操作针对包含 marker 的部分
                    t_el = r.find(qn("w:t"))
                    run_text = t_el.text or ""
                # 2) 在 marker 末尾切，前半=marker 加粗 sz=24，后半保持默认
                if len(run_text) > len(marker):
                    _split_run_at(r, len(marker))
                _run_set_font(r, size_pt=12, bold=True)
                break
        fixed += 1
    return fixed


# ---------------------------------------------------------------------------
# 4.3 引用上标：所有 HYPERLINK 字段码内的 run 设为 superscript
# ---------------------------------------------------------------------------


def _convert_element_hyperlinks_to_fields(doc: Document, *, superscript: bool = True) -> int:
    """把 <w:hyperlink w:anchor="X"> 元素整体替换为 HYPERLINK 字段码。

    某些 Word 版本（尤其 Word for Mac、WPS）对元素式内部锚点跳转兼容性差；
    字段码 HYPERLINK \\l "X" 形式更稳。可选地把整个字段码序列设为上标。
    """
    body = doc.element.body
    anchor_attr = f"{{{W_NS}}}anchor"
    hyperlinks = body.findall(f".//{{{W_NS}}}hyperlink")
    converted = 0
    for hl in hyperlinks:
        anchor = hl.get(anchor_attr)
        if not anchor:
            continue

        runs = list(hl.findall(f"{{{W_NS}}}r"))
        if not runs:
            continue

        replacements: list[OxmlElement] = []

        def fld_run(kind: str) -> OxmlElement:
            r = OxmlElement("w:r")
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), kind)
            r.append(fc)
            return r

        replacements.append(fld_run("begin"))

        r_instr = OxmlElement("w:r")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f' HYPERLINK \\l "{anchor}" '
        r_instr.append(instr)
        replacements.append(r_instr)

        replacements.append(fld_run("separate"))

        for r in runs:
            replacements.append(r)

        replacements.append(fld_run("end"))

        if superscript:
            for r in replacements:
                _run_set_font(r, superscript=True)

        parent = hl.getparent()
        idx = list(parent).index(hl)
        parent.remove(hl)
        for i, elem in enumerate(replacements):
            parent.insert(idx + i, elem)
        converted += 1
    return converted


# ---------------------------------------------------------------------------
# 4.4 表格美化：booktabs 风格 + 整表居中 + 表头加粗
# ---------------------------------------------------------------------------


def _border_el(tag: str, *, sz: int | None) -> OxmlElement:
    el = OxmlElement(tag)
    if sz is None:
        el.set(qn("w:val"), "nil")
    else:
        el.set(qn("w:val"), "single")
        el.set(qn("w:color"), "auto")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
    return el


def _set_cell_borders(tc, *, top: int | None, bottom: int | None) -> None:
    """对 tc（w:tc）设置 booktabs 风格上下边框；左右始终 nil。"""
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    else:
        for child in list(tcBorders):
            tcBorders.remove(child)
    tcBorders.append(_border_el("w:top", sz=top))
    tcBorders.append(_border_el("w:left", sz=None))
    tcBorders.append(_border_el("w:bottom", sz=bottom))
    tcBorders.append(_border_el("w:right", sz=None))


def _beautify_tables(doc: Document) -> int:
    """所有表格：整表居中 + booktabs 三道线 + 表头加粗 + 单元格内文本居中。"""
    body = doc.element.body
    count = 0
    for tbl in body.iter(f"{{{W_NS}}}tbl"):
        # 1. tblPr 加 jc=center
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        jc = tblPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            tblPr.append(jc)
        jc.set(qn("w:val"), "center")

        rows = tbl.findall(qn("w:tr"))
        if not rows:
            continue
        last_idx = len(rows) - 1

        for ridx, tr in enumerate(rows):
            is_header = ridx == 0
            is_last = ridx == last_idx

            # 行级：表头标记 + 整行居中
            trPr = tr.find(qn("w:trPr"))
            if trPr is None:
                trPr = OxmlElement("w:trPr")
                tr.insert(0, trPr)
            tr_jc = trPr.find(qn("w:jc"))
            if tr_jc is None:
                tr_jc = OxmlElement("w:jc")
                trPr.append(tr_jc)
            tr_jc.set(qn("w:val"), "center")
            if is_header:
                if trPr.find(qn("w:tblHeader")) is None:
                    trPr.append(OxmlElement("w:tblHeader"))

            # 单元格
            for tc in tr.findall(qn("w:tc")):
                # 边框
                top_sz = 12 if is_header else None
                bottom_sz = 4 if is_header else (12 if is_last else None)
                _set_cell_borders(tc, top=top_sz, bottom=bottom_sz)

                # 单元格内段落：居中 + 无缩进；表头 run 加粗
                for p in tc.findall(qn("w:p")):
                    pPr = p.find(qn("w:pPr"))
                    if pPr is None:
                        pPr = OxmlElement("w:pPr")
                        p.insert(0, pPr)
                    p_jc = pPr.find(qn("w:jc"))
                    if p_jc is None:
                        p_jc = OxmlElement("w:jc")
                        pPr.append(p_jc)
                    p_jc.set(qn("w:val"), "center")
                    ind = pPr.find(qn("w:ind"))
                    if ind is None:
                        ind = OxmlElement("w:ind")
                        pPr.append(ind)
                    for k in ("w:left", "w:right", "w:firstLine"):
                        ind.set(qn(k), "0")

                    if is_header:
                        for r in p.findall(qn("w:r")):
                            _run_set_font(r, bold=True)

                # 单元格垂直居中
                tcPr = tc.find(qn("w:tcPr"))
                vAlign = tcPr.find(qn("w:vAlign"))
                if vAlign is None:
                    vAlign = OxmlElement("w:vAlign")
                    tcPr.append(vAlign)
                vAlign.set(qn("w:val"), "center")
        count += 1
    return count


# ---------------------------------------------------------------------------
# 4.5 数字直立 / OMML
# ---------------------------------------------------------------------------


def _upright_math_numbers(doc: Document) -> int:
    """OMML 中纯数字 <m:r> 设为 plain（避免被 Word 默认斜体）。"""
    body = doc.element.body
    m_r_tag = f"{{{M_NS}}}r"
    m_t_tag = f"{{{M_NS}}}t"
    m_rpr_tag = f"{{{M_NS}}}rPr"
    m_sty_tag = f"{{{M_NS}}}sty"

    fixed = 0
    for mr in body.iter(m_r_tag):
        joined = "".join((t.text or "") for t in mr.iter(m_t_tag)).strip()
        if not joined or not DIGIT_RE.match(joined):
            continue
        rpr = mr.find(m_rpr_tag)
        if rpr is None:
            rpr = OxmlElement("m:rPr")
            mr.insert(0, rpr)
        sty = rpr.find(m_sty_tag)
        if sty is None:
            sty = OxmlElement("m:sty")
            rpr.insert(0, sty)
        sty.set(qn("m:val"), "p")
        fixed += 1
    return fixed


# ---------------------------------------------------------------------------
# 4.6 目录/题名等自定义样式段落的处理（不依赖 reference.docx 中是否定义该样式）
# ---------------------------------------------------------------------------


def _style_custom_blocks(doc: Document) -> tuple[int, int, int]:
    """渲染 pandoc fenced div 的 custom-style 段落。

    - TOCTitle: 居中、三号黑体加粗，但 *不是* Heading（不进 TOC）
    - ThesisTitleCN: 居中、小二号华文中宋加粗（中文题名）
    - ThesisTitleEN: 居中、小二号 Times New Roman 加粗（英文题名）
    """
    toc_title = title_cn = title_en = 0
    for para in doc.paragraphs:
        sn = para.style.name
        # 设为 Normal 避免污染 outline
        if sn in {"TOCTitle", "ThesisTitleCN", "ThesisTitleEN"}:
            try:
                para.style = doc.styles["Normal"]
            except KeyError:
                pass
        if sn == "TOCTitle":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para._element.findall(qn("w:r")):
                _run_set_font(r, eastasia="黑体", ascii_="Times New Roman",
                              size_pt=16, bold=True, color="000000")
            toc_title += 1
        elif sn == "ThesisTitleCN":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para._element.findall(qn("w:r")):
                _run_set_font(r, eastasia="华文中宋", ascii_="Times New Roman",
                              size_pt=18, bold=True, color="000000")
            title_cn += 1
        elif sn == "ThesisTitleEN":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para._element.findall(qn("w:r")):
                _run_set_font(r, eastasia="Times New Roman",
                              ascii_="Times New Roman",
                              size_pt=18, bold=True, color="000000")
            title_en += 1
    return toc_title, title_cn, title_en


# ===========================================================================
# 后处理总入口
# ===========================================================================


def post_process_docx(path: Path) -> None:
    doc = Document(str(path))

    toc_inserted = False
    bookmarks_added = 0
    centered = 0
    next_bm_id = 10000

    for para in list(doc.paragraphs):
        text = para.text.strip()

        # TOC 字段：在「目录」/「目  录」标题段之后插入
        if not toc_inserted:
            if (
                text.replace(" ", "").replace("\u3000", "") == "目录"
                and para.style.name in {"TOCTitle", "Heading 1"}
            ):
                toc_para = _make_toc_paragraph()
                para._element.addnext(toc_para)
                toc_inserted = True
                continue

        # 参考文献条目：开头形如 "[N] ..."，补插 bookmark name=ref_N
        m = REF_LINE_RE.match(text)
        if m:
            _add_bookmark(para, f"ref_{m.group(1)}", next_bm_id)
            next_bm_id += 1
            bookmarks_added += 1

        # 真图/表 caption 段或含图片段：居中（"图 N"开头的正文引用句不动）
        if _is_real_caption(para) is not None or _paragraph_has_image(para):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            centered += 1

    # Heading 1/2/3 样式（去缩进 + 宋体黑色；H2 15pt / H3 14pt 加粗）
    h1_fixed, h2_fixed, h3_fixed = _restyle_chapter_headings(doc)

    # 参考文献/致谢 H1 采用深大专用样式（覆盖通用 H1）
    refs_h, acks_h = _restyle_special_headings(doc)

    # 参考文献条目：悬挂缩进 + 段前后 0.5 行 + 小五
    refs_entries = _restyle_reference_entries(doc)

    # 隐藏【摘要】/【Abstract】Heading 1（页面不显示，但 TOC 仍抓）
    hidden = _hide_marker_headings(doc)

    # 每个 Heading 1 前 page break（章节另起一页；含【摘要】）
    pbreaks = _add_page_breaks_before_h1(doc)

    # 目录用罗马数字、正文从【摘要】起改阿拉伯
    sect_changes = _setup_section_pagenum(doc)

    # 自定义样式段落（目录标题/中英文题名）
    toc_t, ttl_cn, ttl_en = _style_custom_blocks(doc)

    # 摘要 / Abstract 引导段
    abstract_fixed = _restyle_abstract_paragraphs(doc)

    # 图/表 caption 字号 9pt 小五 + 居中
    captions_fixed = _restyle_captions(doc)

    # 图片宽度统一限制（cx ≤ DEFAULT_MAX_PIC_CX）
    pics_fixed = _resize_images(doc)

    # 表格美化（booktabs + 居中 + 表头加粗）
    tables_fixed = _beautify_tables(doc)

    # 元素式 hyperlink → 字段码（同时上标化引用）
    hl_converted = _convert_element_hyperlinks_to_fields(doc, superscript=True)

    # OMML 数字直立
    upright = _upright_math_numbers(doc)

    doc.save(str(path))
    print(
        f"[post] toc={toc_inserted} bookmarks={bookmarks_added} centered={centered}\n"
        f"       h1={h1_fixed} h2={h2_fixed} h3={h3_fixed} refs_h={refs_h} "
        f"acks_h={acks_h} hidden_h1={hidden} pbreaks={pbreaks} sect={sect_changes}\n"
        f"       toc_title={toc_t} title_cn={ttl_cn} title_en={ttl_en} "
        f"abstract={abstract_fixed} refs_entries={refs_entries}\n"
        f"       captions={captions_fixed} pics_resized={pics_fixed} "
        f"tables={tables_fixed} hyperlinks={hl_converted} math_nums={upright}"
    )


# ===========================================================================
# main
# ===========================================================================


def main() -> None:
    merged = merge_markdown()
    print(f"[merge] -> {merged}")
    out = convert_to_docx(merged)
    print(f"[docx]  -> {out}")
    post_process_docx(out)
    print(f"[done]  -> {out}")


if __name__ == "__main__":
    main()
